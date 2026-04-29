"""
resume_parser.py — Multi-format resume parsing utility.

Supported formats:
  • PDF  (.pdf)  — text extraction via PyPDF2; OCR fallback via pytesseract
  • DOCX (.docx) — python-docx
  • Images (.png, .jpg, .jpeg, .bmp, .tiff, .webp) — pytesseract OCR

Install all dependencies:
    pip install PyPDF2 python-docx pytesseract Pillow
    # Also install Tesseract OCR engine on your OS:
    #   Windows: https://github.com/UB-Mannheim/tesseract/wiki
    #   Linux:   sudo apt install tesseract-ocr
    #   macOS:   brew install tesseract
"""

from __future__ import annotations
import re
import io

# ── Optional imports with graceful fallbacks ─────────────────────
try:
    import fitz
    _Fitz = True
except ImportError:
    _Fitz = False

try:
    import PyPDF2
    _PDF = True
except ImportError:
    _PDF = False

try:
    import docx as _docx_lib
    _DOCX = True
except ImportError:
    _DOCX = False

try:
    from PIL import Image
    import pytesseract
    _OCR = True
except ImportError:
    _OCR = False

# ── Skill detection patterns ─────────────────────────────────────
SKILL_MAP = {
    # Programming languages
    r"\bpython\b":          "Python",
    r"\bjava\b":            "Java",
    r"c\+\+":               "C++",
    r"\bc#\b":              "C#",
    r"\bjavascript\b":      "JavaScript",
    r"\btypescript\b":      "TypeScript",
    r"\bkotlin\b":          "Kotlin",
    r"\bswift\b":           "Swift",
    r"\bruntime\b|\brust\b": "Rust",
    r"\bgo(?:lang)?\b":     "Go",
    r"\bphp\b":             "PHP",
    r"\bruby\b":            "Ruby",
    r"\bscala\b":           "Scala",
    r"\br\b(?= programming| language| studio)": "R",

    # Web / Frontend
    r"\breact\.?js\b|\breact\b": "React",
    r"\bnode\.?js\b":       "Node.js",
    r"\bvue\.?js\b":        "Vue.js",
    r"\bangular\b":         "Angular",
    r"\bnext\.?js\b":       "Next.js",
    r"\bhtml5?\b":          "HTML",
    r"\bcss3?\b":           "CSS",
    r"\bbootstrap\b":       "Bootstrap",
    r"\btailwind\b":        "Tailwind CSS",

    # Backend / Frameworks
    r"\bdjango\b":          "Django",
    r"\bflask\b":           "Flask",
    r"\bfastapi\b":         "FastAPI",
    r"\bspring boot\b|\bspring\b": "Spring Boot",
    r"\blaravel\b":         "Laravel",
    r"\bexpress\.?js\b":    "Express.js",

    # Data / AI / ML
    r"machine learning":    "Machine Learning",
    r"deep learning":       "Deep Learning",
    r"data science":        "Data Science",
    r"data analysis":       "Data Analysis",
    r"\btensorflow\b":      "TensorFlow",
    r"\bpytorch\b":         "PyTorch",
    r"\bkeras\b":           "Keras",
    r"\bnlp\b|natural language processing": "NLP",
    r"computer vision":     "Computer Vision",
    r"\bpandas\b":          "Pandas",
    r"\bnumpy\b":           "NumPy",
    r"\bscikit.?learn\b":   "Scikit-Learn",
    r"\bopencv\b":          "OpenCV",
    r"\bmatplotlib\b":      "Matplotlib",
    r"\bseaborn\b":         "Seaborn",
    r"\bpower\s?bi\b":      "Power BI",
    r"\btableau\b":         "Tableau",
    r"\bexcel\b":           "Excel",

    # Databases
    r"\bsql\b":             "SQL",
    r"\bmysql\b":           "MySQL",
    r"\bpostgresql\b|\bpostgres\b": "PostgreSQL",
    r"\bmongodb\b":         "MongoDB",
    r"\bredis\b":           "Redis",
    r"\belasticsearch\b":   "Elasticsearch",
    r"\bhadoop\b":          "Hadoop",
    r"\bspark\b":           "Apache Spark",
    r"\bcassandra\b":       "Cassandra",
    r"\bfirebase\b":        "Firebase",
    r"\bsqlite\b":          "SQLite",
    r"\boracle\b":          "Oracle DB",

    # Cloud / DevOps
    r"\baws\b|amazon web services": "AWS",
    r"\bgcp\b|google cloud":        "Google Cloud",
    r"\bazure\b":           "Azure",
    r"\bdocker\b":          "Docker",
    r"\bkubernetes\b|\bk8s\b": "Kubernetes",
    r"\bterraform\b":       "Terraform",
    r"\bjenkins\b":         "Jenkins",
    r"\bgithub actions\b":  "GitHub Actions",
    r"\bgit\b":             "Git",
    r"\blinux\b|\bubuntu\b": "Linux",
    r"\bci/cd\b|\bcicd\b":  "CI/CD",
    r"\bnginx\b":           "Nginx",

    # Engineering / Domain
    r"\bembedded\b":        "Embedded Systems",
    r"\bvlsi\b":            "VLSI",
    r"\bmatlab\b":          "MATLAB",
    r"\bansys\b":           "ANSYS",
    r"\bautocad\b":         "AutoCAD",
    r"\bsolidworks\b":      "SolidWorks",
    r"\bcatia\b":           "CATIA",
    r"\bplc\b":             "PLC/SCADA",
    r"\biot\b":             "IoT",
    r"\barduino\b":         "Arduino",
    r"\braspberry pi\b":    "Raspberry Pi",

    # Security / Blockchain
    r"\bcybersecurity\b|ethical hack|\bpenetration test": "Cybersecurity",
    r"\bblockchain\b":      "Blockchain",
    r"\bweb3\b":            "Web3",

    # Soft skills (detected contextually)
    r"\bteam\s?work\b|\bcollaboration\b": "Teamwork",
    r"\bleadership\b":      "Leadership",
    r"\bcommunication\b":   "Communication",
    r"\bproblem.?solving\b": "Problem Solving",
    r"\bagile\b|\bscrum\b": "Agile/Scrum",
}

# ── Keyword lists for counting ───────────────────────────────────
_INTERN_KWS = [
    r"\binternship\b", r"\bintern\b", r"\btrainee\b",
    r"\bsummer\s+(?:project|training|intern)\b",
    r"\bindustry training\b", r"\bwork experience\b",
]
_PROJECT_KWS = [
    r"\bproject\b", r"\bbuilt\b", r"\bdeveloped\b",
    r"\bimplemented\b", r"\bdesigned\b", r"\bcreated\b",
    r"\bdeploy(?:ed)?\b", r"\bengineered\b",
]
_CERT_KWS = [
    r"\bcertif", r"\bcoursera\b", r"\budemy\b", r"\bnptel\b",
    r"\bedx\b", r"\bmooc\b", r"\blinkedin learning\b",
    r"\baws certified\b", r"\bgoogle certified\b",
    r"\bmicrosoft certified\b", r"\bcisco\b",
]
_CGPA_PATTERN = r"(?:cgpa|gpa|cpi)\s*[:\-]?\s*(\d+(?:\.\d+)?)"
_EMAIL_PATTERN = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
_PHONE_PATTERN = r"(?:\+91[\s\-]?)?[6-9]\d{9}"


# ══════════════════════════════════════════════════════════════════
# TEXT EXTRACTION FUNCTIONS
# ══════════════════════════════════════════════════════════════════

def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF first, then fallback to PDF text or OCR."""
    if _Fitz:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages_text = [page.get_text("text") or "" for page in doc]
            full_text = "\n".join(pages_text)
            if len(full_text.strip()) < 120 and _OCR:
                return _ocr_pdf(file_bytes)
            return full_text
        except Exception:
            pass

    if _PDF:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                t = page.extract_text() or ""
                pages_text.append(t)
            full_text = "\n".join(pages_text)
            if len(full_text.strip()) < 120 and _OCR:
                return _ocr_pdf(file_bytes)
            return full_text
        except Exception:
            pass

    return _ocr_pdf(file_bytes) if _OCR else ""


def _ocr_pdf(file_bytes: bytes) -> str:
    """Convert PDF pages to images and OCR them."""
    if not _OCR:
        return ""
    try:
        # Try pdf2image if available
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(file_bytes, dpi=200)
        except Exception:
            return ""
        texts = []
        for img in images:
            t = pytesseract.image_to_string(img, config="--psm 6")
            texts.append(t)
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    if not _DOCX:
        return ""
    try:
        doc = _docx_lib.Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_from_image(file_bytes: bytes) -> str:
    """OCR an image file (PNG, JPG, BMP, TIFF, WEBP)."""
    if not _OCR:
        return ""
    try:
        img = Image.open(io.BytesIO(file_bytes))
        # Pre-process: convert to RGB if needed
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img, config="--psm 6")
        return text
    except Exception:
        return ""


def extract_text(uploaded_file) -> tuple[str, str]:
    """
    Main entry point. Accepts a Streamlit UploadedFile.
    Returns (text: str, method: str) where method describes what was used.
    """
    name = uploaded_file.name.lower()
    raw  = uploaded_file.read()

    if name.endswith(".pdf"):
        text = _extract_from_pdf(raw)
        method = "PDF text extraction" + (" + OCR fallback" if not text else "")
    elif name.endswith(".docx"):
        text = _extract_from_docx(raw)
        method = "DOCX text extraction"
    elif name.endswith(".doc"):
        # .doc (legacy Word) — try docx parser anyway; often fails
        text = _extract_from_docx(raw)
        method = "Legacy DOC (partial support)"
    elif any(name.endswith(ext) for ext in (".png",".jpg",".jpeg",".bmp",".tiff",".tif",".webp")):
        text = _extract_from_image(raw)
        method = "OCR (image)"
    else:
        text = ""
        method = "unsupported format"

    return text.strip(), method


# ══════════════════════════════════════════════════════════════════
# PARSING FUNCTIONS
# ══════════════════════════════════════════════════════════════════

def detect_skills(text: str) -> list[str]:
    t = text.lower()
    return [label for pat, label in SKILL_MAP.items() if re.search(pat, t, re.IGNORECASE)]


def detect_cgpa(text: str) -> float | None:
    """Try to extract CGPA/GPA from resume text."""
    m = re.search(_CGPA_PATTERN, text, re.IGNORECASE)
    if m:
        val = float(m.group(1))
        if 0.0 < val <= 10.0:
            return round(val, 2)
        if 0.0 < val <= 4.0:
            return round(val * 2.5, 2)   # convert 4-point scale
    return None


def _count_kws(text: str, kws: list[str], cap: int = 5) -> int:
    t = text.lower()
    return min(sum(1 for kw in kws if re.search(kw, t, re.IGNORECASE)), cap)


def parse_resume(text: str) -> dict:
    """
    Full resume parse. Returns structured dict.
    """
    if not text:
        return {
            "skills": [], "internships": 0, "projects": 0,
            "certifications": 0, "cgpa": None,
            "has_email": False, "has_phone": False,
            "word_count": 0,
        }

    return {
        "skills":        detect_skills(text),
        "internships":   _count_kws(text, _INTERN_KWS, 5),
        "projects":      min(_count_kws(text, _PROJECT_KWS, 12), 10),
        "certifications":_count_kws(text, _CERT_KWS, 8),
        "cgpa":          detect_cgpa(text),
        "has_email":     bool(re.search(_EMAIL_PATTERN, text)),
        "has_phone":     bool(re.search(_PHONE_PATTERN, text)),
        "word_count":    len(text.split()),
    }


def get_support_status() -> dict:
    """Return which parsers are available."""
    return {
        "PDF":  _PDF,
        "DOCX": _DOCX,
        "OCR (images / scanned PDFs)": _OCR,
    }


SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"]
STREAMLIT_TYPES      = ["pdf", "docx", "doc", "png", "jpg", "jpeg", "bmp", "tiff", "webp"]
